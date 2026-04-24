from io import BytesIO

import pytest
from docx import Document as DocxDocument

from semantic_reasoning_agent.core.container import get_app_container
from semantic_reasoning_agent.persistence.models import (
    OntologyBuildORM,
    OntologyBuildStepORM,
)
from semantic_reasoning_agent.persistence.models import OntologyEntityFactORM, OntologyRelationFactORM
from semantic_reasoning_agent.schemas.ontology import OntologyBuildCreateRequest
from semantic_reasoning_agent.services.ontology_service import (
    OntologyBuildError,
    OntologyBuildNotFoundError,
    OntologyPublishError,
)
from semantic_reasoning_agent.domain.ontology.models import (
    ExtractedEntity,
    ExtractedRelation,
    ExtractionResult,
    OntologyNarrative,
)
from semantic_reasoning_agent.schemas.ontology import (
    OntologyDraftPublishRequest,
    OntologyGraphDraftNodeRequest,
    OntologyGraphDraftRelationRequest,
)


class _RateLimitedExtractor:
    def classify_document_domain(self, document) -> str:  # noqa: ANN001
        return "pending"

    def extract_ontology_candidates(  # noqa: ANN001
        self,
        document,
        workspace_id=None,
        provider=None,
        model=None,
    ):
        raise _FakeRateLimitError(
            "raw upstream error",
            body={
                "error": {
                    "message": "Provider returned error",
                    "metadata": {
                        "raw": (
                            "qwen/qwen3-next-80b-a3b-instruct:free is temporarily "
                            "rate-limited upstream. Please retry shortly."
                        )
                    },
                }
            },
        )

    def summarize_ontology(  # noqa: ANN001
        self,
        document,
        *,
        workspace_id=None,
        provider=None,
        model=None,
        domain=None,
    ) -> OntologyNarrative:
        return OntologyNarrative(title="Rate Limited Build", summary="Pending ontology summary.")


class _FakeRateLimitError(Exception):
    def __init__(self, message: str, *, body: dict) -> None:
        super().__init__(message)
        self.status_code = 429
        self.body = body


class _UnavailableExtractor:
    def classify_document_domain(self, document) -> str:  # noqa: ANN001
        return "pending"

    def extract_ontology_candidates(  # noqa: ANN001
        self,
        document,
        workspace_id=None,
        provider=None,
        model=None,
    ):
        from semantic_reasoning_agent.domain.ontology.models import ExtractionResult

        return ExtractionResult(domain="unavailable", entities=[], relations=[])

    def summarize_ontology(  # noqa: ANN001
        self,
        document,
        *,
        workspace_id=None,
        provider=None,
        model=None,
        domain=None,
    ) -> OntologyNarrative:
        return OntologyNarrative(title="Unavailable Ontology", summary="Model is not ready.")


class _ChunkedTraceExtractor:
    def classify_document_domain(self, document) -> str:  # noqa: ANN001
        return "pending"

    def extract_ontology_candidates(  # noqa: ANN001
        self,
        document,
        workspace_id=None,
        provider=None,
        model=None,
    ) -> ExtractionResult:
        base_provenance = {
            "extractor": "fake",
            "provider": provider or "openrouter",
            "model": model or "test-model",
            "prompt_version": "v2",
            "source_document_id": document.document_id,
        }
        entities = [
            ExtractedEntity(
                name="Alpha",
                canonical_name="Alpha",
                resolution_key="alpha",
                entity_type="system",
                confidence=0.9,
                source_chunk_id=None,
                evidence_text="Alpha evidence.",
                provenance=base_provenance,
                aliases={"A"},
            ),
            ExtractedEntity(
                name="Beta",
                canonical_name="Beta",
                resolution_key="beta",
                entity_type="system",
                confidence=0.8,
                source_chunk_id=None,
                evidence_text="Beta evidence.",
                provenance=base_provenance,
                aliases=set(),
            ),
        ]
        relations = [
            ExtractedRelation(
                source_resolution_key="alpha",
                target_resolution_key="beta",
                source_name="Alpha",
                target_name="Beta",
                relation_type="depends_on",
                confidence=0.7,
                source_chunk_id=None,
                evidence_text="Alpha depends on Beta.",
                provenance=base_provenance,
            )
        ]
        return ExtractionResult(
            domain="chunked_ops",
            entities=entities,
            relations=relations,
            trace={
                "chunks": [
                    {"stage": "entities", "chunk_index": 0},
                    {"stage": "relations", "chunk_index": 0},
                    {"stage": "entities", "chunk_index": 1},
                ],
                "errors": ["chunk[1] entities: empty_payload"],
            },
        )

    def summarize_ontology(  # noqa: ANN001
        self,
        document,
        *,
        workspace_id=None,
        provider=None,
        model=None,
        domain=None,
    ) -> OntologyNarrative:
        return OntologyNarrative(title="Chunked Trace Build", summary="Chunked extraction traces are present.")


class _EmptyExtractionExtractor:
    def classify_document_domain(self, document) -> str:  # noqa: ANN001
        return "pending"

    def extract_ontology_candidates(  # noqa: ANN001
        self,
        document,
        workspace_id=None,
        provider=None,
        model=None,
    ) -> ExtractionResult:
        return ExtractionResult(domain="general", entities=[], relations=[], trace={"chunks": [], "errors": []})

    def summarize_ontology(  # noqa: ANN001
        self,
        document,
        *,
        workspace_id=None,
        provider=None,
        model=None,
        domain=None,
    ) -> OntologyNarrative:
        return OntologyNarrative(title="Empty Ontology", summary="No entities were extracted.")


class _MergeExtractor:
    def classify_document_domain(self, document) -> str:  # noqa: ANN001
        return "pending"

    def extract_ontology_candidates(  # noqa: ANN001
        self,
        document,
        workspace_id=None,
        provider=None,
        model=None,
    ) -> ExtractionResult:
        base_provenance = {
            "extractor": "merge-test",
            "provider": provider or "openrouter",
            "model": model or "test-model",
            "source_document_id": document.document_id,
        }
        entities = [
            ExtractedEntity(
                name="Alpha Platform",
                canonical_name="Alpha Platform",
                resolution_key="alpha",
                entity_type="system",
                confidence=0.95,
                source_chunk_id=None,
                evidence_text="Alpha platform evolved from Alpha.",
                provenance=base_provenance,
                aliases={"Alpha"},
            ),
            ExtractedEntity(
                name="Gamma Service",
                canonical_name="Gamma Service",
                resolution_key="gamma-service",
                entity_type="service",
                confidence=0.9,
                source_chunk_id=None,
                evidence_text="Gamma service was introduced.",
                provenance=base_provenance,
                aliases={"Gamma"},
            ),
        ]
        relations = [
            ExtractedRelation(
                source_resolution_key="alpha",
                target_resolution_key="gamma-service",
                source_name="Alpha Platform",
                target_name="Gamma Service",
                relation_type="integrates_with",
                confidence=0.88,
                source_chunk_id=None,
                evidence_text="Alpha integrates with Gamma.",
                provenance=base_provenance,
            )
        ]
        return ExtractionResult(
            domain="delivery",
            entities=entities,
            relations=relations,
            trace={"chunks": [{"stage": "entities", "chunk_index": 0}], "errors": []},
        )

    def summarize_ontology(  # noqa: ANN001
        self,
        document,
        *,
        workspace_id=None,
        provider=None,
        model=None,
        domain=None,
    ) -> OntologyNarrative:
        return OntologyNarrative(title="Merged Ontology", summary="Merged ontology publish.")


def test_process_build_marks_rate_limited_extraction_as_failed(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    ontology_service._ontology_extractor = _RateLimitedExtractor()

    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="openrouter",
            extraction_model="qwen/qwen3-next-80b-a3b-instruct:free",
        )
    )

    with pytest.raises(OntologyBuildError, match="upstream rate limit"):
        ontology_service.process_build(build.id)

    failed_build = ontology_service.get_build(build.id)
    assert failed_build.status.value == "failed"
    assert failed_build.domain is None
    assert "temporarily rate-limited upstream" in (failed_build.error_message or "")


def test_delete_build_removes_failed_build_and_steps(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    ontology_service._ontology_extractor = _RateLimitedExtractor()

    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="openrouter",
            extraction_model="qwen/qwen3-next-80b-a3b-instruct:free",
        )
    )

    with pytest.raises(OntologyBuildError):
        ontology_service.process_build(build.id)

    ontology_service.delete_build(build.id)

    with pytest.raises(OntologyBuildNotFoundError):
        ontology_service.get_build(build.id)

    with get_app_container().database_manager.session() as session:
        assert session.get(OntologyBuildORM, build.id) is None
        assert (
            session.query(OntologyBuildStepORM)
            .filter(OntologyBuildStepORM.build_id == build.id)
            .count()
            == 0
        )


def test_delete_build_rejects_non_failed_build(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )

    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="echo",
            extraction_model="local-echo",
        )
    )

    with pytest.raises(OntologyBuildError, match="Only failed ontology builds can be deleted."):
        ontology_service.delete_build(build.id)

    reloaded_build = ontology_service.get_build(build.id)
    assert reloaded_build.status.value == "pending"


def test_process_build_fails_when_extraction_model_is_unavailable(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    ontology_service._ontology_extractor = _UnavailableExtractor()

    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="openai",
            extraction_model="gpt-5-mini",
        )
    )

    with pytest.raises(OntologyBuildError, match="model is unavailable"):
        ontology_service.process_build(build.id)

    failed_build = ontology_service.get_build(build.id)
    assert failed_build.status.value == "failed"
    assert "unavailable" in (failed_build.error_message or "").lower()


def test_process_build_uses_extraction_domain_when_classification_is_deferred(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )

    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="echo",
            extraction_model="local-echo",
        )
    )
    ontology_service.process_build(build.id)

    completed_build = ontology_service.get_build(build.id)
    assert completed_build.status.value == "completed"
    assert completed_build.domain == "test_domain"
    assert completed_build.ontology_title == "Delivery Dependencies"
    assert completed_build.ontology_summary
    extract_entities_step = next(
        step for step in completed_build.steps if step.name == "extract_entities"
    )
    extract_relations_step = next(
        step for step in completed_build.steps if step.name == "extract_relations"
    )
    entities_trace = extract_entities_step.metadata.get("safe_trace")
    relations_trace = extract_relations_step.metadata.get("safe_trace")
    assert isinstance(entities_trace, dict)
    assert entities_trace.get("domain") == "test_domain"
    assert entities_trace.get("entity_count") == completed_build.entity_count
    assert isinstance(relations_trace, dict)
    assert relations_trace.get("domain") == "test_domain"
    assert relations_trace.get("relation_count") == completed_build.relation_count
    resolve_step = next(step for step in completed_build.steps if step.name == "resolve_entities")
    resolve_meta = resolve_step.metadata or {}
    merge_stats = resolve_meta.get("merge_stats")
    assert isinstance(merge_stats, dict)
    assert merge_stats.get("canonical_entities_after_merge") == completed_build.entity_count
    assert merge_stats.get("canonical_relations_after_merge") == completed_build.relation_count
    assert completed_build.pending_entity_count == 0
    assert completed_build.pending_relation_count == 0


def test_process_build_propagates_chunked_trace_metadata(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    ontology_service._ontology_extractor = _ChunkedTraceExtractor()

    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="openrouter",
            extraction_model="test-model",
        )
    )
    ontology_service.process_build(build.id)
    completed_build = ontology_service.get_build(build.id)
    assert completed_build.status.value == "completed"
    assert completed_build.domain == "chunked_ops"

    extract_entities_step = next(
        step for step in completed_build.steps if step.name == "extract_entities"
    )
    extract_relations_step = next(
        step for step in completed_build.steps if step.name == "extract_relations"
    )

    entities_trace = extract_entities_step.metadata.get("safe_trace")
    relations_trace = extract_relations_step.metadata.get("safe_trace")
    assert isinstance(entities_trace, dict)
    assert entities_trace.get("chunk_count") == 3
    assert entities_trace.get("errors") == ["chunk[1] entities: empty_payload"]
    assert isinstance(relations_trace, dict)
    assert relations_trace.get("chunk_count") == 3
    assert relations_trace.get("errors") == ["chunk[1] entities: empty_payload"]


def test_graph_draft_round_trip_publish_and_reset(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    document = document_service.upload_document(
        filename="draft.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="echo",
            extraction_model="local-echo",
        )
    )
    ontology_service.process_build(build.id)
    draft = ontology_service.create_draft_node(
        OntologyGraphDraftNodeRequest(name="Alpha", entity_type="initiative")
    )
    existing_node = next(entity for entity in draft.entities if entity.name == "Alpha")
    draft = ontology_service.create_draft_node(
        OntologyGraphDraftNodeRequest(name="Control Tower", entity_type="capability")
    )
    added_node = next(entity for entity in draft.entities if entity.name == "Control Tower")
    draft = ontology_service.create_draft_relation(
        OntologyGraphDraftRelationRequest(
            source_entity_id=existing_node.id,
            target_entity_id=added_node.id,
            relation_type="monitors",
            evidence_text="Draft editor link.",
        )
    )
    assert draft.has_changes is True
    assert any(relation.relation_type == "monitors" for relation in draft.relations)

    publish = ontology_service.publish_graph_draft(
        OntologyDraftPublishRequest(build_id=build.id)
    )
    assert publish.version.version_number == 1
    graph = ontology_service.get_graph()
    assert any(entity.name == "Control Tower" for entity in graph.entities)
    assert any(relation.relation_type == "monitors" for relation in graph.relations)
    with get_app_container().database_manager.session() as session:
        assert session.query(OntologyEntityFactORM).count() == 0
        assert session.query(OntologyRelationFactORM).count() == 0

    reset = ontology_service.reset_graph_draft()
    assert reset.has_changes is False


def test_publish_graph_draft_from_completed_build_without_manual_draft_changes(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    document = document_service.upload_document(
        filename="publish-from-build.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="echo",
            extraction_model="local-echo",
        )
    )
    ontology_service.process_build(build.id)

    completed_build = ontology_service.get_build(build.id)
    assert completed_build.status.value == "completed"
    assert completed_build.has_publishable_entities is True

    publish = ontology_service.publish_graph_draft(OntologyDraftPublishRequest(build_id=build.id))
    assert publish.version.version_number == 1
    assert publish.build.status.value == "published"

    graph = ontology_service.get_graph()
    assert len(graph.entities) >= 1
    assert len(graph.relations) >= 1


def test_publish_graph_draft_fails_when_build_has_no_publishable_entities(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    ontology_service._ontology_extractor = _EmptyExtractionExtractor()
    document = document_service.upload_document(
        filename="publish-empty-build.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="openrouter",
            extraction_model="test-model",
        )
    )
    ontology_service.process_build(build.id)

    completed_build = ontology_service.get_build(build.id)
    assert completed_build.status.value == "completed"
    assert completed_build.has_publishable_entities is False

    with pytest.raises(OntologyPublishError, match="No ontology entities are available to publish."):
        ontology_service.publish_graph_draft(OntologyDraftPublishRequest(build_id=build.id))


def test_publish_graph_draft_merges_existing_version_with_new_build_seed(
    document_service,
    ontology_service,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )

    first_document = document_service.upload_document(
        filename="publish-merge-v1.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    first_build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=first_document.id,
            extraction_provider="echo",
            extraction_model="local-echo",
        )
    )
    ontology_service.process_build(first_build.id)
    ontology_service.publish_graph_draft(OntologyDraftPublishRequest(build_id=first_build.id))

    ontology_service._ontology_extractor = _MergeExtractor()
    second_document = document_service.upload_document(
        filename="publish-merge-v2.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    second_build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=second_document.id,
            extraction_provider="openrouter",
            extraction_model="test-model",
        )
    )
    ontology_service.process_build(second_build.id)
    publish = ontology_service.publish_graph_draft(OntologyDraftPublishRequest(build_id=second_build.id))
    assert publish.version.version_number == 2

    graph = ontology_service.get_graph()
    assert any(entity.resolution_key == "alpha" and entity.name == "Alpha Platform" for entity in graph.entities)
    assert any(entity.resolution_key == "gamma-service" for entity in graph.entities)
    assert any(relation.relation_type == "integrates_with" for relation in graph.relations)


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Delivery Plan", level=1)
    document.add_paragraph("Alpha initiative depends on the beta system for approvals.")
    document.add_paragraph("Beta system uses Audit service.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
