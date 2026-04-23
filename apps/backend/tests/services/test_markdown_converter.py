from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

from semantic_reasoning_agent.documents.chunking import MarkdownChunker
from semantic_reasoning_agent.documents.converters import MarkdownConverterService


def test_native_converter_extracts_docx_to_markdown() -> None:
    converter = MarkdownConverterService()
    payload = converter.convert("source.docx", _build_docx_bytes("Alpha depends on Beta"))

    assert payload.document_type == "docx"
    assert payload.converter_name == "native"
    assert "Alpha depends on Beta" in payload.markdown


def test_native_converter_extracts_csv_to_markdown_and_chunker_splits() -> None:
    converter = MarkdownConverterService()
    chunker = MarkdownChunker(chunk_size=120, overlap=20)

    payload = converter.convert(
        "table.csv",
        b"building_id,name,floors\nB001,Tower A,12\nB002,Tower B,8\n",
    )
    chunks = chunker.split(payload.markdown)

    assert payload.document_type == "csv"
    assert payload.converter_name == "native"
    assert "building_id" in payload.markdown.lower()
    assert len(chunks) >= 1
    assert chunks[0].text


def test_supported_types_include_markitdown_all_extensions() -> None:
    converter = MarkdownConverterService()
    supported = set(converter.supported_types())

    assert {"xls", "msg", "wav", "mp3", "m4a", "epub", "zip"}.issubset(supported)


def _build_docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_heading("Delivery Plan", level=1)
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
