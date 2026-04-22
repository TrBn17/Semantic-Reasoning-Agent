from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from semantic_reasoning_agent.main import app


client = TestClient(app)


def _build_docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_heading("Knowledge", level=1)
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_agent_capability_catalog_and_tool_view_are_available() -> None:
    catalog_response = client.get("/api/v1/agent-capabilities/catalog")
    tools_response = client.get("/api/v1/agent-capabilities/tools")

    assert catalog_response.status_code == 200
    assert tools_response.status_code == 200

    catalog = catalog_response.json()
    assert any(item["preset"] == "internal_qa" for item in catalog["presets"])
    assert "retrieval" in catalog["tool_families"]

    tools = tools_response.json()
    assert any(item["tool_name"] == "retrieval.internal" for item in tools)


def test_agent_profile_crud_returns_typed_capability_fields() -> None:
    create_response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": "workspace-demo",
            "name": "Ontology Curator",
            "capability_preset": "ontology_curator",
            "tool_policy": {"mode": "preset", "blocked_tools": ["ontology.lookup"]},
            "knowledge_pack_ids": ["pack-a", "pack-b"],
            "evidence_policy": {
                "allowed_sources": ["internal_chunk"],
                "allow_model_only_fallback": False,
            },
            "task_models": [
                {"task_type": "chat", "provider": "echo", "model": "local-echo"}
            ],
        },
    )

    assert create_response.status_code == 201
    profile = create_response.json()
    assert profile["capability_preset"] == "ontology_curator"
    assert profile["tool_policy"]["blocked_tools"] == ["ontology.lookup"]
    assert profile["knowledge_pack_ids"] == ["pack-a", "pack-b"]
    assert profile["evidence_policy"]["allow_model_only_fallback"] is False
    assert profile["policy_config"]["capability_preset"] == "ontology_curator"

    patch_response = client.patch(
        f"/api/v1/agents/profiles/{profile['id']}",
        json={
            "capability_preset": "graph_explorer",
            "knowledge_pack_ids": ["pack-c"],
        },
    )
    assert patch_response.status_code == 200
    updated = patch_response.json()
    assert updated["capability_preset"] == "graph_explorer"
    assert updated["knowledge_pack_ids"] == ["pack-c"]


def test_knowledge_pack_crud_round_trips_document_membership() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "pack-source.docx",
                _build_docx_bytes("Pack scoped content."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document = upload_response.json()

    create_response = client.post(
        "/api/v1/knowledge-packs",
        json={
            "workspace_id": document["workspace_id"],
            "name": "Core Pack",
            "description": "Scoped docs",
            "document_ids": [document["id"]],
        },
    )
    assert create_response.status_code == 201
    pack = create_response.json()
    assert pack["document_ids"] == [document["id"]]

    list_response = client.get("/api/v1/knowledge-packs", params={"workspace_id": document["workspace_id"]})
    assert list_response.status_code == 200
    assert any(item["id"] == pack["id"] for item in list_response.json())

    patch_response = client.patch(
        f"/api/v1/knowledge-packs/{pack['id']}",
        json={"description": "Updated scoped docs"},
    )
    assert patch_response.status_code == 200
    assert patch_response.json()["description"] == "Updated scoped docs"


def test_empty_knowledge_scope_uses_model_fallback_only_when_policy_allows() -> None:
    allow_response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": "workspace-demo",
            "name": "Fallback Agent",
            "capability_preset": "internal_qa",
            "knowledge_pack_ids": [],
            "evidence_policy": {"allow_model_only_fallback": True},
            "task_models": [
                {"task_type": "chat", "provider": "echo", "model": "local-echo"}
            ],
        },
    )
    deny_response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": "workspace-demo",
            "name": "No Fallback Agent",
            "capability_preset": "internal_qa",
            "knowledge_pack_ids": [],
            "evidence_policy": {"allow_model_only_fallback": False},
            "task_models": [
                {"task_type": "chat", "provider": "echo", "model": "local-echo"}
            ],
        },
    )

    allow_profile = allow_response.json()
    deny_profile = deny_response.json()

    allow_conversation = client.post(
        "/api/v1/conversations",
        json={
            "title": "Fallback conversation",
            "workspace_id": "workspace-demo",
            "agent_profile_id": allow_profile["id"],
        },
    ).json()
    deny_conversation = client.post(
        "/api/v1/conversations",
        json={
            "title": "No fallback conversation",
            "workspace_id": "workspace-demo",
            "agent_profile_id": deny_profile["id"],
        },
    ).json()

    allow_chat = client.post(
        "/api/v1/chat/messages",
        json={"conversation_id": allow_conversation["id"], "content": "hello"},
    )
    deny_chat = client.post(
        "/api/v1/chat/messages",
        json={"conversation_id": deny_conversation["id"], "content": "hello"},
    )

    assert allow_chat.status_code == 201
    assert deny_chat.status_code == 201
    assert allow_chat.json()["reply"]["content"] == "echo: hello"
    assert deny_chat.json()["reply"]["content"] == "No indexed document or graph context matched that question."


def test_agent_pack_scopes_chat_retrieval_to_pack_documents() -> None:
    allowed_upload = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "allowed.docx",
                _build_docx_bytes("Alpha initiative depends on Beta system."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"title": "Allowed Knowledge"},
    )
    blocked_upload = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "blocked.docx",
                _build_docx_bytes("Gamma initiative depends on Delta service."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"title": "Blocked Knowledge"},
    )
    allowed_doc = allowed_upload.json()
    blocked_doc = blocked_upload.json()

    pack_response = client.post(
        "/api/v1/knowledge-packs",
        json={
            "workspace_id": allowed_doc["workspace_id"],
            "name": "Allowed Pack",
            "document_ids": [allowed_doc["id"]],
        },
    )
    pack = pack_response.json()

    profile_response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": allowed_doc["workspace_id"],
            "name": "Pack Agent",
            "capability_preset": "internal_qa",
            "knowledge_pack_ids": [pack["id"]],
            "evidence_policy": {"allow_model_only_fallback": False},
            "task_models": [
                {"task_type": "chat", "provider": "echo", "model": "local-echo"}
            ],
        },
    )
    profile = profile_response.json()
    conversation = client.post(
        "/api/v1/conversations",
        json={
            "title": "Scoped chat",
            "workspace_id": allowed_doc["workspace_id"],
            "agent_profile_id": profile["id"],
        },
    ).json()

    chat_response = client.post(
        "/api/v1/chat/messages",
        json={
            "conversation_id": conversation["id"],
            "content": "What depends on the Beta system?",
            "document_ids": [allowed_doc["id"], blocked_doc["id"]],
            "top_k": 3,
        },
    )

    assert chat_response.status_code == 201
    payload = chat_response.json()
    assert payload["citations"]
    assert {item["document_id"] for item in payload["citations"]} == {allowed_doc["id"]}
    assert "Allowed Knowledge" in payload["reply"]["content"]


def test_document_override_cannot_escape_agent_pack_scope() -> None:
    allowed_upload = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "allowed-only.docx",
                _build_docx_bytes("This text is allowed."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"title": "Allowed Only"},
    )
    blocked_upload = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "blocked-only.docx",
                _build_docx_bytes("Blocked secret content."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"title": "Blocked Only"},
    )
    allowed_doc = allowed_upload.json()
    blocked_doc = blocked_upload.json()

    pack = client.post(
        "/api/v1/knowledge-packs",
        json={
            "workspace_id": allowed_doc["workspace_id"],
            "name": "Allowed Only Pack",
            "document_ids": [allowed_doc["id"]],
        },
    ).json()
    profile = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": allowed_doc["workspace_id"],
            "name": "Strict Pack Agent",
            "knowledge_pack_ids": [pack["id"]],
            "evidence_policy": {"allow_model_only_fallback": False},
            "task_models": [
                {"task_type": "chat", "provider": "echo", "model": "local-echo"}
            ],
        },
    ).json()
    conversation = client.post(
        "/api/v1/conversations",
        json={
            "title": "Strict scope",
            "workspace_id": allowed_doc["workspace_id"],
            "agent_profile_id": profile["id"],
        },
    ).json()

    chat_response = client.post(
        "/api/v1/chat/messages",
        json={
            "conversation_id": conversation["id"],
            "content": "Blocked secret content",
            "document_ids": [blocked_doc["id"]],
            "top_k": 3,
        },
    )

    assert chat_response.status_code == 201
    payload = chat_response.json()
    assert payload["citations"] == []
    assert payload["reply"]["content"] == "No indexed document or graph context matched that question."
