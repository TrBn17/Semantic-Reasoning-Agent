from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from semantic_reasoning_agent.main import app


client = TestClient(app)


def test_list_workflows_includes_task_resolve_chat() -> None:
    response = client.get("/api/v1/workflows")

    assert response.status_code == 200
    workflows = response.json()
    assert any(item["workflow_id"] == "task.resolve.chat" for item in workflows)


def test_openapi_marks_task_resolve_public_and_workflow_run_internal() -> None:
    openapi = app.openapi()

    task_resolve = openapi["paths"]["/api/v1/tasks/resolve"]["post"]
    workflow_run = openapi["paths"]["/api/v1/workflows/{workflow_id}/run"]["post"]

    assert task_resolve["x-api-audience"] == "public"
    assert task_resolve["x-api-primary"] is True
    assert workflow_run["x-api-audience"] == "internal/admin"
    assert workflow_run["x-api-primary"] is False
    assert "tasks/resolve" in workflow_run["description"]


def test_resolve_task_returns_retrieval_context_and_evidence() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "workflow-source.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document = upload_response.json()

    response = client.post(
        "/api/v1/tasks/resolve",
        json={
            "content": "What depends on the beta system?",
            "use_retrieval": True,
            "document_ids": [document["id"]],
            "top_k": 2,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert "Relevant context:" in payload["content"]
    assert len(payload["citations"]) >= 1


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Delivery Plan", level=1)
    document.add_paragraph("Alpha initiative depends on the beta system for approvals.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
