from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from semantic_reasoning_agent.documents.parsers.pdf_marker_parser import PdfMarkerParser
from semantic_reasoning_agent.main import app


client = TestClient(app)


def test_upload_docx_indexes_document_and_reports_jobs() -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "plan.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"title": "Execution Plan"},
    )

    assert response.status_code == 201
    document = response.json()
    assert document["status"] == "indexed"
    assert document["document_type"] == "docx"
    assert document["chunk_count"] >= 1

    jobs_response = client.get(f"/api/v1/documents/{document['id']}/jobs")
    assert jobs_response.status_code == 200
    jobs = jobs_response.json()
    assert [job["name"] for job in jobs] == [
        "parse_document",
        "build_chunks",
        "embed_chunks",
        "upsert_qdrant",
    ]
    assert all(job["status"] == "completed" for job in jobs)


def test_pdf_search_returns_page_citation() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("brief.pdf", _build_pdf_bytes("Alpha program requires approval on page one."), "application/pdf")},
    )
    document = upload_response.json()

    search_response = client.post(
        "/api/v1/retrieval/search",
        json={"query": "Which program requires approval?", "document_ids": [document["id"]], "top_k": 2},
    )

    assert search_response.status_code == 200
    results = search_response.json()["results"]
    assert len(results) >= 1
    assert results[0]["citation"]["page_number"] == 1
    assert results[0]["citation"]["location_label"] == "page 1"


def test_xlsx_search_returns_sheet_and_row_range_citation() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("sales.xlsx", _build_xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    document = upload_response.json()

    search_response = client.post(
        "/api/v1/retrieval/search",
        json={"query": "What revenue did East have in January?", "document_ids": [document["id"]], "top_k": 3},
    )

    assert search_response.status_code == 200
    results = search_response.json()["results"]
    assert len(results) >= 1
    assert results[0]["citation"]["sheet_name"] == "Sales"
    assert "Sales rows" in results[0]["citation"]["location_label"]


def test_csv_search_returns_row_range_citation() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("sales.csv", _build_csv_bytes(), "text/csv")},
    )
    document = upload_response.json()

    search_response = client.post(
        "/api/v1/retrieval/search",
        json={"query": "What revenue did East have in January?", "document_ids": [document["id"]], "top_k": 3},
    )

    assert search_response.status_code == 200
    results = search_response.json()["results"]
    assert len(results) >= 1
    assert results[0]["citation"]["sheet_name"] is None
    assert "rows" in results[0]["citation"]["location_label"]


def test_chat_with_retrieval_returns_citations() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("plan.docx", _build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    document = upload_response.json()

    conversation_response = client.post(
        "/api/v1/conversations",
        json={"title": "Grounded chat", "provider": "echo", "model": "local-echo"},
    )
    conversation = conversation_response.json()

    chat_response = client.post(
        "/api/v1/chat/messages",
        json={
            "conversation_id": conversation["id"],
            "content": "What depends on the beta system?",
            "provider": "echo",
            "model": "local-echo",
            "use_retrieval": True,
            "document_ids": [document["id"]],
            "top_k": 2,
        },
    )

    assert chat_response.status_code == 201
    payload = chat_response.json()
    assert len(payload["citations"]) >= 1
    assert "Relevant context:" in payload["reply"]["content"]
    assert payload["citations"][0]["document_type"] == "docx"


def test_document_reprocess_and_reindex_endpoint() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("sales.xlsx", _build_xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    document = upload_response.json()

    reprocess_response = client.post(f"/api/v1/documents/{document['id']}/reprocess")
    assert reprocess_response.status_code == 200
    assert reprocess_response.json()["document"]["status"] == "indexed"

    reindex_response = client.post("/api/v1/retrieval/reindex", json={"document_ids": [document["id"]]})
    assert reindex_response.status_code == 200
    assert reindex_response.json()["reindexed_document_ids"] == [document["id"]]


def test_upload_queues_ingestion_when_dispatcher_does_not_execute(document_service) -> None:
    queued_document_ids: list[str] = []

    class FakeDispatcher:
        def enqueue_document_processing(self, document_id: str) -> None:
            queued_document_ids.append(document_id)

    original_dispatcher = document_service._task_dispatcher
    document_service._task_dispatcher = FakeDispatcher()
    try:
        response = client.post(
            "/api/v1/documents/upload",
            files={
                "file": (
                    "plan.docx",
                    _build_docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    finally:
        document_service._task_dispatcher = original_dispatcher

    assert response.status_code == 201
    document = response.json()
    assert document["status"] == "uploaded"
    assert document["chunk_count"] == 0
    assert queued_document_ids == [document["id"]]

    jobs_response = client.get(f"/api/v1/documents/{document['id']}/jobs")
    jobs = jobs_response.json()
    assert all(job["status"] == "pending" for job in jobs)


def test_upload_pdf_accurate_mode_persists_ingestion_options(monkeypatch) -> None:
    def fake_render(self, content: bytes, options):  # noqa: ANN001
        del self, content

        class _Page:
            def __init__(self, children):
                self.children = children

        class _Block:
            def __init__(self, markdown: str):
                self.markdown = markdown

        class _Rendered:
            def __init__(self):
                self.children = [_Page([_Block("Accurate OCR text")])]

        assert options.pdf_mode == "accurate"
        return _Rendered(), "Accurate OCR text"

    monkeypatch.setattr(PdfMarkerParser, "_render_pdf", fake_render)

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("accurate.pdf", _build_pdf_bytes("ignored"), "application/pdf")},
        data={"pdf_mode": "accurate"},
    )

    assert upload_response.status_code == 201
    document = upload_response.json()
    assert document["ingestion_options"]["pdf_mode"] == "accurate"
    assert document["parser_version"] == "marker-v1"

    reprocess_response = client.post(f"/api/v1/documents/{document['id']}/reprocess")
    assert reprocess_response.status_code == 200
    assert reprocess_response.json()["document"]["ingestion_options"]["pdf_mode"] == "accurate"


def test_upload_pdf_with_invalid_mode_returns_400() -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("bad.pdf", _build_pdf_bytes("ignored"), "application/pdf")},
        data={"pdf_mode": "turbo"},
    )

    assert response.status_code == 400
    assert "pdf_mode" in response.json()["detail"]


def _build_pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, text)
    pdf.save()
    return buffer.getvalue()


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Delivery Plan", level=1)
    document.add_paragraph("Alpha initiative depends on the beta system for approvals.")
    document.add_heading("Dependencies", level=2)
    document.add_paragraph("The beta system also depends on the audit service.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sales"
    sheet.append(["Month", "Region", "Revenue", "Product"])
    sheet.append(["January", "East", 120, "Alpha"])
    sheet.append(["January", "West", 90, "Beta"])
    sheet.append(["February", "East", 150, "Alpha"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _build_csv_bytes() -> bytes:
    return "\n".join(
        [
            "Month,Region,Revenue,Product",
            "January,East,120,Alpha",
            "January,West,90,Beta",
            "February,East,150,Alpha",
        ]
    ).encode("utf-8")
