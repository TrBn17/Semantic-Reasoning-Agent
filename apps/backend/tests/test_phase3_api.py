from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from semantic_reasoning_agent.main import app
client = TestClient(app)


def test_ontology_build_review_and_publish_flow() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "ontology-source.docx",
                _build_ontology_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document = upload_response.json()

    build_response = client.post(
        "/api/v1/ontology/builds",
        json={"document_id": document["id"], "provider": "echo", "model": "local-echo"},
    )
    assert build_response.status_code == 201
    build = build_response.json()
    assert build["status"] == "completed"
    assert build["provider"] == "echo"
    assert build["model"] == "local-echo"
    assert build["relation_count"] >= 2
    assert build["pending_entity_count"] == build["entity_count"]
    assert build["pending_relation_count"] == build["relation_count"]
    assert {entity_type["name"] for entity_type in build["entity_type_definitions"]} >= {"test_thing"}
    assert {relation_type["name"] for relation_type in build["relation_type_definitions"]} >= {
        "depends_on",
        "uses",
    }
    assert [step["name"] for step in build["steps"]] == [
        "classify_document_domain",
        "extract_entities",
        "extract_relations",
        "resolve_entities",
        "build_graph_upsert_plan",
        "sync_neo4j",
    ]

    entities_response = client.get(
        f"/api/v1/ontology/builds/{build['id']}/entities",
        params={"review_status": "pending_review"},
    )
    assert entities_response.status_code == 200
    entities = entities_response.json()
    assert {entity["resolution_key"] for entity in entities} >= {
        "alpha-initiative",
        "beta-system",
        "audit-service",
    }
    assert all("run_id" in entity["provenance"] for entity in entities)
    assert all(entity["provenance"]["prompt_version"] == "v1" for entity in entities)

    relations_response = client.get(f"/api/v1/ontology/builds/{build['id']}/relations")
    assert relations_response.status_code == 200
    relations = relations_response.json()
    assert {
        (relation["relation_type"], relation["source_name"], relation["target_name"])
        for relation in relations
    } >= {
        ("depends_on", "Alpha Initiative", "Beta System"),
        ("uses", "Beta System", "Audit Service"),
    }

    blocked_publish_response = client.post(f"/api/v1/ontology/builds/{build['id']}/publish")
    assert blocked_publish_response.status_code == 400
    assert "no approved entities" in blocked_publish_response.json()["detail"].lower()

    for entity in entities:
        review_response = client.post(
            f"/api/v1/ontology/entities/{entity['id']}/review",
            json={"action": "approve"},
        )
        assert review_response.status_code == 200
        assert review_response.json()["status"] == "approved"

    for relation in relations:
        review_response = client.post(
            f"/api/v1/ontology/relations/{relation['id']}/review",
            json={"action": "approve"},
        )
        assert review_response.status_code == 200
        assert review_response.json()["status"] == "approved"

    publish_response = client.post(f"/api/v1/ontology/builds/{build['id']}/publish")
    assert publish_response.status_code == 200
    publish_payload = publish_response.json()
    assert publish_payload["build"]["status"] == "published"
    assert publish_payload["version"]["version_number"] == 1
    assert publish_payload["version"]["entity_type_count"] >= 1
    assert publish_payload["version"]["relation_type_count"] >= 2
    assert publish_payload["version"]["relation_count"] >= 2

    graph_response = client.get("/api/v1/ontology/graph")
    assert graph_response.status_code == 200
    graph = graph_response.json()
    assert graph["version"]["source_build_id"] == build["id"]
    assert {entity_type["name"] for entity_type in graph["entity_type_definitions"]} >= {"test_thing"}
    assert {relation_type["name"] for relation_type in graph["relation_type_definitions"]} >= {
        "depends_on",
        "uses",
    }
    assert {entity["resolution_key"] for entity in graph["entities"]} >= {
        "alpha-initiative",
        "beta-system",
        "audit-service",
    }
    assert {relation["relation_type"] for relation in graph["relations"]} >= {
        "depends_on",
        "uses",
    }


def test_ontology_build_queues_when_dispatcher_does_not_execute(ontology_service) -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "queued-ontology.docx",
                _build_ontology_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document = upload_response.json()

    queued_build_ids: list[str] = []

    class FakeDispatcher:
        def enqueue_ontology_build_processing(self, build_id: str) -> None:
            queued_build_ids.append(build_id)

    original_dispatcher = ontology_service._task_dispatcher
    ontology_service._task_dispatcher = FakeDispatcher()
    try:
        build_response = client.post(
            "/api/v1/ontology/builds",
            json={"document_id": document["id"]},
        )
    finally:
        ontology_service._task_dispatcher = original_dispatcher

    assert build_response.status_code == 201
    build = build_response.json()
    assert build["status"] == "pending"
    assert build["entity_count"] == 0
    assert build["relation_count"] == 0
    assert build["entity_type_definitions"] == []
    assert build["relation_type_definitions"] == []
    assert queued_build_ids == [build["id"]]
    assert all(step["status"] == "pending" for step in build["steps"])


def test_publish_records_snapshot_for_graphiti_publish(ontology_service) -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "neo4j-publish.docx",
                _build_ontology_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    document = upload_response.json()

    build_response = client.post("/api/v1/ontology/builds", json={"document_id": document["id"]})
    build = build_response.json()
    entities = client.get(f"/api/v1/ontology/builds/{build['id']}/entities").json()
    relations = client.get(f"/api/v1/ontology/builds/{build['id']}/relations").json()

    for entity in entities:
        client.post(f"/api/v1/ontology/entities/{entity['id']}/review", json={"action": "approve"})
    for relation in relations:
        client.post(f"/api/v1/ontology/relations/{relation['id']}/review", json={"action": "approve"})

    class FakePublisher:
        def __init__(self) -> None:
            self.snapshot = None

        def is_enabled(self) -> bool:
            return True

        def publish(self, snapshot=None) -> None:
            self.snapshot = snapshot

    fake_publisher = FakePublisher()
    original_publisher = ontology_service._graph_publisher
    ontology_service._graph_publisher = fake_publisher
    try:
        publish_response = client.post(f"/api/v1/ontology/builds/{build['id']}/publish")
        graph_response = client.get("/api/v1/ontology/graph")
    finally:
        ontology_service._graph_publisher = original_publisher

    assert publish_response.status_code == 200
    assert fake_publisher.snapshot is not None
    assert fake_publisher.snapshot.version.source_build_id == build["id"]
    assert {entity_type.name for entity_type in fake_publisher.snapshot.entity_type_definitions} >= {
        "test_thing",
    }
    assert graph_response.status_code == 200
    graph = graph_response.json()
    assert graph["version"]["source_build_id"] == build["id"]
    assert {relation["relation_type"] for relation in graph["relations"]} >= {
        "depends_on",
        "uses",
    }


def _build_ontology_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Alpha initiative depends on Beta system for approvals.")
    document.add_paragraph("Beta system uses Audit service.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
