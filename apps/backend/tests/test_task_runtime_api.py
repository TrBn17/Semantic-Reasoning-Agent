from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from semantic_reasoning_agent.main import app


client = TestClient(app)


def test_workflows_endpoint_lists_registered_runtime_workflows() -> None:
    response = client.get("/api/v1/workflows")
    assert response.status_code == 200
    body = response.json()
    workflow_ids = {item["workflow_id"] for item in body}
    assert {
        "answer_resolution",
        "document_ingestion",
        "ontology_build",
        "ontology_design",
        "ontology_candidate_build",
        "review_publish",
    } <= workflow_ids


def test_tasks_resolve_persists_task_and_tool_calls_for_grounded_answer() -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "plan.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document = upload_response.json()

    response = client.post(
        "/api/v1/tasks/resolve",
        json={
            "entrypoint": "chat",
            "content": "What depends on the beta system?",
            "workspace_id": document["workspace_id"],
            "provider": "echo",
            "model": "local-echo",
            "use_retrieval": True,
            "document_ids": [document["id"]],
            "top_k": 2,
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["workflow_id"] == "answer_resolution"
    assert body["status"] == "completed"
    assert "Relevant context:" in body["reply"]
    assert len(body["citations"]) >= 1
    assert any(item["tool_name"] == "retrieval.internal" for item in body["tool_calls"])

    task_response = client.get(f"/api/v1/tasks/{body['task_id']}")
    assert task_response.status_code == 200
    assert task_response.json()["workflow_id"] == "answer_resolution"

    tool_calls_response = client.get(f"/api/v1/tasks/{body['task_id']}/tool-calls")
    assert tool_calls_response.status_code == 200
    tool_calls = tool_calls_response.json()
    assert any(item["tool_name"] == "retrieval.internal" for item in tool_calls)


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Delivery Plan", level=1)
    document.add_paragraph("Alpha initiative depends on the beta system for approvals.")
    document.add_heading("Dependencies", level=2)
    document.add_paragraph("The beta system also depends on the audit service.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
