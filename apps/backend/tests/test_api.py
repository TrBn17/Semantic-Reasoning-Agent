from io import BytesIO

from fastapi.testclient import TestClient
from docx import Document as DocxDocument
import pytest

from semantic_reasoning_agent.main import app
from semantic_reasoning_agent.schemas.ontology import OntologyBuildCreateRequest
from semantic_reasoning_agent.services.ontology_service import OntologyBuildError
from semantic_reasoning_agent.services.provider_models_service import (
    AnthropicModelsClient,
    OpenAIModelsClient,
    OpenRouterModelsClient,
)


client = TestClient(app)


def test_healthcheck() -> None:
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_delete_failed_ontology_build_via_api(
    document_service,
    ontology_service,
    monkeypatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )

    class _DeleteApiRateLimitedExtractor:
        def classify_document_domain(self, document) -> str:  # noqa: ANN001
            return "pending"

        def extract_ontology_candidates(self, document, workspace_id=None, provider=None, model=None):  # noqa: ANN001
            raise RuntimeError("simulated extraction failure")

        def summarize_ontology(self, document, *, workspace_id=None, provider=None, model=None, domain=None):  # noqa: ANN001
            raise AssertionError("summarize_ontology should not be called after extraction failure")

    ontology_service._ontology_extractor = _DeleteApiRateLimitedExtractor()
    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="openrouter",
            extraction_model="test-model",
        )
    )

    with pytest.raises(OntologyBuildError):
        ontology_service.process_build(build.id)

    response = client.delete(f"/api/v1/ontology/builds/{build.id}")
    assert response.status_code == 204

    missing = client.get(f"/api/v1/ontology/builds/{build.id}")
    assert missing.status_code == 404


def test_delete_pending_ontology_build_via_api_returns_bad_request(
    document_service,
    ontology_service,
    monkeypatch,
) -> None:
    document = document_service.upload_document(
        filename="ontology-source.docx",
        content=_build_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    monkeypatch.setattr(
        ontology_service._task_dispatcher,
        "enqueue_ontology_build_processing",
        lambda build_id: None,
    )
    build = ontology_service.create_build(
        OntologyBuildCreateRequest(
            document_id=document.id,
            extraction_provider="echo",
            extraction_model="local-echo",
        )
    )

    response = client.delete(f"/api/v1/ontology/builds/{build.id}")
    assert response.status_code == 400
    assert response.json()["detail"] == "Only failed ontology builds can be deleted."


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Delivery Plan", level=1)
    document.add_paragraph("Alpha initiative depends on the beta system for approvals.")
    document.add_paragraph("Beta system uses Audit service.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_auth_me_contract() -> None:
    response = client.get("/api/v1/auth/me")
    assert response.status_code == 200

    payload = response.json()
    assert payload["id"] == "user-demo"
    assert payload["active_workspace"]["id"] == "workspace-demo"


def test_model_catalog_includes_local_echo() -> None:
    response = client.get("/api/v1/settings/models")
    assert response.status_code == 200

    models = response.json()
    assert any(item["provider"] == "echo" and item["ready"] is True for item in models)


def test_public_settings_bootstrap_exposes_business_first_contract() -> None:
    response = client.get("/api/v1/settings")

    assert response.status_code == 200
    payload = response.json()
    assert payload["workspace"]["id"] == "workspace-demo"
    assert payload["workspace"]["name"] == "Demo Workspace"
    assert "model_catalog" in payload
    assert "task_defaults" in payload
    assert "models" not in payload
    assert any(item["use_case"] == "chat_default" for item in payload["task_defaults"])
    assert any(item["provider"] == "echo" for item in payload["model_catalog"])


def test_removed_compatibility_aliases_return_not_found() -> None:
    canonical = client.get("/api/v1/settings/models")
    alias_models = client.get("/api/v1/models")
    alias_catalog = client.get("/api/v1/agents/catalog")
    alias_settings_get = client.get("/api/v1/agents/settings")
    alias_settings_put = client.put(
        "/api/v1/agents/settings",
        json={"workspace_id": "workspace-demo", "providers": [], "task_assignments": []},
    )

    assert canonical.status_code == 200
    assert alias_models.status_code == 404
    assert alias_catalog.status_code == 404
    assert alias_settings_get.status_code == 404
    assert alias_settings_put.status_code == 404


def test_public_settings_update_round_trips_with_canonical_contract(monkeypatch) -> None:
    async def _raise_runtime_error(self):
        raise RuntimeError("upstream list-models unavailable")

    monkeypatch.setattr(OpenRouterModelsClient, "get_models", _raise_runtime_error)
    update_response = client.put(
        "/api/v1/settings",
        json={
            "workspace_id": "workspace-settings",
            "providers": [
                {
                    "provider": "openrouter",
                    "enabled": True,
                    "values": {"OPENROUTER_API_KEY": "demo-openrouter-key"},
                }
            ],
            "task_defaults": [
                {
                    "use_case": "chat_default",
                    "provider": "openrouter",
                    "model": "nvidia/nemotron-3-super-120b-a12b:free",
                }
            ],
        },
    )

    assert update_response.status_code == 200
    payload = update_response.json()
    assert payload["workspace"]["id"] == "workspace-settings"
    assert payload["preferred_default_chat_model"]["provider"] == "openrouter"

    settings_response = client.get("/api/v1/settings", params={"workspace_id": "workspace-settings"})
    assert settings_response.status_code == 200
    assert settings_response.json()["workspace"]["id"] == "workspace-settings"
    chat_assignment = next(
        item for item in settings_response.json()["task_defaults"] if item["task_type"] == "chat"
    )
    assert chat_assignment["provider"] == "openrouter"
    assert chat_assignment["model"] == "nvidia/nemotron-3-super-120b-a12b:free"


def test_settings_can_persist_provider_and_task_assignments() -> None:
    update_response = client.put(
        "/api/v1/settings",
        json={
            "workspace_id": "workspace-demo",
            "providers": [
                {
                    "provider": "openai",
                    "enabled": True,
                    "values": {"OPENAI_API_KEY": "demo-openai-key"},
                }
            ],
            "task_defaults": [
                {
                    "use_case": "chat_default",
                    "provider": "openai",
                    "model": "gpt-5-mini",
                }
            ],
        },
    )
    assert update_response.status_code == 200
    payload = update_response.json()

    openai_provider = next(item for item in payload["providers"] if item["provider"] == "openai")
    assert openai_provider["values"][0]["configured"] is True
    assert openai_provider["values"][0]["source"] == "database"
    assert "demo-openai-key" not in openai_provider["values"][0]["display_value"]

    chat_assignment = next(item for item in payload["task_defaults"] if item["task_type"] == "chat")
    assert chat_assignment["provider"] == "openai"
    assert chat_assignment["model"] == "gpt-5-mini"
    assert chat_assignment["ready"] is False
    assert "runtime" in chat_assignment["reason"].lower()


def test_model_catalog_uses_workspace_specific_provider_config() -> None:
    client.put(
        "/api/v1/settings",
        json={
            "workspace_id": "workspace-openai",
            "providers": [
                {
                    "provider": "openai",
                    "enabled": True,
                    "values": {"OPENAI_API_KEY": "demo-openai-key"},
                }
            ],
            "task_defaults": [],
        },
    )

    response = client.get("/api/v1/settings", params={"workspace_id": "workspace-openai"})
    assert response.status_code == 200
    payload = response.json()
    openai = next(item for item in payload["providers"] if item["provider"] == "openai")
    assert openai["enabled"] is True
    assert openai["ready"] is False


def test_provider_models_endpoint_uses_workspace_specific_provider_config(monkeypatch) -> None:
    async def _raise_runtime_error(self):
        raise RuntimeError("upstream list-models unavailable")

    monkeypatch.setattr(OpenAIModelsClient, "get_models", _raise_runtime_error)

    client.put(
        "/api/v1/settings",
        json={
            "workspace_id": "workspace-openai-provider-endpoint",
            "providers": [
                {
                    "provider": "openai",
                    "enabled": True,
                    "values": {"OPENAI_API_KEY": "demo-openai-key"},
                }
            ],
            "task_defaults": [],
        },
    )

    response = client.get(
        "/api/v1/providers/openai/models",
        params={"workspace_id": "workspace-openai-provider-endpoint"},
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["provider"] == "openai"
    assert any(item["id"] == "gpt-5-mini" for item in payload["models"])


def test_model_catalog_falls_back_to_maintained_anthropic_catalog(monkeypatch) -> None:
    async def _raise_runtime_error(self):
        raise RuntimeError("upstream list-models unavailable")

    monkeypatch.setattr(AnthropicModelsClient, "get_models", _raise_runtime_error)

    update_response = client.put(
        "/api/v1/settings",
        json={
            "workspace_id": "workspace-anthropic",
            "providers": [
                {
                    "provider": "anthropic",
                    "enabled": True,
                    "values": {
                        "ANTHROPIC_API_KEY": "demo-ant-key",
                        "ANTHROPIC_BASE_URL": "https://example.invalid",
                    },
                }
            ],
            "task_defaults": [],
        },
    )
    assert update_response.status_code == 200

    response = client.get("/api/v1/settings/models", params={"workspace_id": "workspace-anthropic"})
    assert response.status_code == 200
    models = response.json()

    anthropic_models = [item for item in models if item["provider"] == "anthropic"]
    assert anthropic_models
    assert any(item["model"] == "claude-sonnet-4-5" for item in anthropic_models)
    assert any(item["model"] == "claude-opus-4-0" and item["label"] == "Claude Opus 4.0" for item in anthropic_models)
    assert any(item["model"] == "claude-sonnet-4-5" and item["label"] == "Claude Sonnet 4.5" for item in anthropic_models)


def test_create_conversation_and_send_echo_message() -> None:
    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Phase 1 prep", "provider": "echo", "model": "local-echo"},
    )
    assert create_response.status_code == 201

    conversation = create_response.json()
    send_response = client.post(
        "/api/v1/chat/messages",
        json={
            "conversation_id": conversation["id"],
            "content": "hello",
            "provider": "echo",
            "model": "local-echo",
        },
    )
    assert send_response.status_code == 201

    payload = send_response.json()
    assert payload["reply"]["role"] == "assistant"
    assert payload["reply"]["content"] == "echo: hello"
    assert payload["citations"] == []
    assert payload["tool_calls"] == []
    assert payload["conversation_id"] == conversation["id"]


def test_chat_stream_endpoint_emits_structured_events() -> None:
    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Streaming chat", "provider": "echo", "model": "local-echo"},
    )
    conversation = create_response.json()

    with client.stream(
        "POST",
        "/api/v1/chat/messages/stream",
        json={
            "conversation_id": conversation["id"],
            "content": "hello stream",
            "provider": "echo",
            "model": "local-echo",
        },
    ) as response:
        assert response.status_code == 200
        payload = "".join(response.iter_text())

    assert "event: message_start" in payload
    assert "event: content_delta" in payload
    assert "event: message_complete" in payload
    assert "echo: hello stream" in payload


def test_default_agent_profile_drives_new_conversation_model() -> None:
    tools_response = client.get("/api/v1/search-tools")
    assert tools_response.status_code == 200
    tools = tools_response.json()
    rag_tool = next(item for item in tools if item["system_key"] == "workspace_default_rag")
    ontology_tool = next(
        item for item in tools if item["system_key"] == "workspace_default_ontology_search"
    )
    profile_response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": "workspace-demo",
            "name": "General Analyst",
            "system_prompt": "You are a precise analyst.",
            "is_default": True,
            "task_models": [
                {
                    "task_type": "chat",
                    "provider": "echo",
                    "model": "local-echo",
                }
            ],
            "tool_assignments": [
                {
                    "slot": "rag",
                    "tool_name": "supersearch.docs",
                    "config_id": rag_tool["id"],
                    "enabled": True,
                    "position": 0,
                },
                {
                    "slot": "ontology_search",
                    "tool_name": "supersearch.graph",
                    "config_id": ontology_tool["id"],
                    "enabled": True,
                    "position": 1,
                },
            ],
        },
    )
    assert profile_response.status_code == 201
    profile = profile_response.json()

    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Profile chat", "workspace_id": "workspace-demo"},
    )
    assert create_response.status_code == 201
    conversation = create_response.json()
    assert conversation["agent_profile_id"] == profile["id"]
    assert conversation["uses_model_override"] is False
    assert conversation["provider"] == "echo"
    assert conversation["effective_tool_names"] == ["supersearch.docs", "supersearch.graph"]
    assert [item["label"] for item in conversation["effective_tool_bindings"]] == [
        rag_tool["name"],
        ontology_tool["name"],
    ]


def test_agent_profile_rejects_slot_assignment_when_tool_policy_blocks_it() -> None:
    tools = client.get("/api/v1/search-tools").json()
    rag_tool = next(item for item in tools if item["system_key"] == "workspace_default_rag")

    response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": "workspace-demo",
            "name": "Blocked slots",
            "tool_policy": {
                "mode": "blocklist",
                "allowed_tools": [],
                "blocked_tools": ["supersearch.docs"],
            },
            "tool_assignments": [
                {
                    "slot": "rag",
                    "tool_name": "supersearch.docs",
                    "config_id": rag_tool["id"],
                    "enabled": True,
                    "position": 0,
                }
            ],
        },
    )

    assert response.status_code == 400
    assert "blocked" in response.json()["detail"].lower()


def test_conversation_model_override_persists_per_session() -> None:
    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Override chat", "workspace_id": "workspace-demo"},
    )
    conversation = create_response.json()

    update_response = client.patch(
        f"/api/v1/conversations/{conversation['id']}/model-selection",
        json={"provider": "echo", "model": "local-echo", "workspace_id": "workspace-demo"},
    )
    assert update_response.status_code == 200
    updated = update_response.json()
    assert updated["uses_model_override"] is True
    assert updated["provider"] == "echo"
    assert updated["model"] == "local-echo"


def test_conversation_model_override_rejects_unready_model() -> None:
    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Bad override chat", "workspace_id": "workspace-demo"},
    )
    conversation = create_response.json()

    update_response = client.patch(
        f"/api/v1/conversations/{conversation['id']}/model-selection",
        json={"provider": "openai", "model": "gpt-5-mini", "workspace_id": "workspace-demo"},
    )
    assert update_response.status_code == 400
    assert "is not ready" in update_response.json()["detail"]


def test_profile_can_block_chat_model_override() -> None:
    profile_response = client.post(
        "/api/v1/agents/profiles",
        json={
            "workspace_id": "workspace-demo",
            "name": "Locked agent",
            "allow_chat_model_override": False,
            "task_models": [
                {
                    "task_type": "chat",
                    "provider": "echo",
                    "model": "local-echo",
                }
            ],
        },
    )
    profile = profile_response.json()

    create_response = client.post(
        "/api/v1/conversations",
        json={
            "title": "Locked conversation",
            "workspace_id": "workspace-demo",
            "agent_profile_id": profile["id"],
        },
    )
    conversation = create_response.json()

    update_response = client.patch(
        f"/api/v1/conversations/{conversation['id']}/model-selection",
        json={"provider": "echo", "model": "local-echo", "workspace_id": "workspace-demo"},
    )
    assert update_response.status_code == 400
    assert "does not allow chat model override" in update_response.json()["detail"]


def test_unready_provider_is_rejected() -> None:
    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Anthropic later", "provider": "anthropic", "model": "claude-sonnet-4-5"},
    )
    conversation = create_response.json()

    send_response = client.post(
        "/api/v1/chat/messages",
        json={
            "conversation_id": conversation["id"],
            "content": "hello",
            "provider": "anthropic",
            "model": "claude-sonnet-4-5",
        },
    )

    assert send_response.status_code == 400
    assert "not ready yet" in send_response.json()["detail"]


def test_conversation_creation_falls_back_to_ready_model_when_assignment_is_blocked() -> None:
    update_response = client.put(
        "/api/v1/agents/settings",
        json={
            "workspace_id": "workspace-demo",
            "providers": [
                {
                    "provider": "openai",
                    "enabled": True,
                    "values": {"OPENAI_API_KEY": "demo-openai-key"},
                }
            ],
            "task_assignments": [
                {
                    "task_type": "chat",
                    "provider": "openai",
                    "model": "gpt-5-mini",
                }
            ],
        },
    )
    assert update_response.status_code == 200

    create_response = client.post(
        "/api/v1/conversations",
        json={"title": "Fallback chat", "workspace_id": "workspace-demo"},
    )
    assert create_response.status_code == 201
    conversation = create_response.json()

    assert conversation["provider"] == "echo"
    assert conversation["model"] == "local-echo"
    assert conversation["uses_model_override"] is False
