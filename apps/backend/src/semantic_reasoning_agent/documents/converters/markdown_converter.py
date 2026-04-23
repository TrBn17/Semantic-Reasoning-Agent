from __future__ import annotations

import csv
from importlib.metadata import PackageNotFoundError, version
from io import BytesIO
from pathlib import Path
from typing import Any
from zipfile import BadZipFile

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from semantic_reasoning_agent.documents.errors import DocumentProcessingError, UnsupportedDocumentTypeError
from semantic_reasoning_agent.documents.models import ConvertedDocument

FALLBACK_CONVERTER_NAME = "markitdown"

_SUPPORTED_TYPES = {
    "pdf",
    "docx",
    "xlsx",
    "xls",
    "ppt",
    "pptx",
    "html",
    "htm",
    "epub",
    "csv",
    "json",
    "xml",
    "md",
    "markdown",
    "txt",
    "jpg",
    "jpeg",
    "png",
    "gif",
    "webp",
    "bmp",
    "tif",
    "tiff",
    "zip",
    "msg",
    "eml",
    "wav",
    "mp3",
    "m4a",
}


class MarkdownConverterService:
    def supports(self, filename: str) -> bool:
        return _document_type(filename) in _SUPPORTED_TYPES

    def supported_types(self) -> tuple[str, ...]:
        return tuple(sorted(_SUPPORTED_TYPES))

    def convert(
        self,
        filename: str,
        content: bytes,
        *,
        title: str | None = None,
        content_type: str | None = None,
    ) -> ConvertedDocument:
        del content_type
        document_type = _document_type(filename)
        if document_type not in _SUPPORTED_TYPES:
            raise UnsupportedDocumentTypeError(
                f"Unsupported document type '{document_type}'. "
                "Supported types: "
                + ", ".join(sorted(_SUPPORTED_TYPES))
                + "."
            )
        if not content:
            raise DocumentProcessingError("Uploaded file is empty.")

        converter_name, converter_version, markdown, metadata = _convert_to_markdown(filename, content)
        markdown = (markdown or "").strip()
        if not markdown:
            raise DocumentProcessingError(f"No extractable text was found in '{filename}'.")

        return ConvertedDocument(
            document_type=document_type,
            title=title or Path(filename).stem,
            markdown=markdown,
            converter_name=converter_name,
            converter_version=converter_version,
            metadata=metadata,
            warnings=(),
        )


def _convert_to_markdown(filename: str, content: bytes) -> tuple[str, str, str, dict[str, Any]]:
    document_type = _document_type(filename)
    native_converter = _NATIVE_CONVERTERS.get(document_type)
    if native_converter is None:
        markdown, metadata = _convert_with_markitdown(filename, content)
        return FALLBACK_CONVERTER_NAME, _converter_version("markitdown"), markdown, metadata

    try:
        markdown, metadata = native_converter(filename, content)
    except (BadZipFile, ValueError) as exc:
        raise DocumentProcessingError(f"Failed to parse '{filename}': {exc}") from exc

    metadata = dict(metadata)
    metadata.setdefault("parser_backend", "native")
    return "native", "builtin", markdown, metadata


def _convert_pdf(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    reader = PdfReader(BytesIO(content))
    lines: list[str] = []
    page_count = len(reader.pages)
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        if lines:
            lines.append("")
        lines.append(f"## Page {index}")
        lines.append(text)
    return "\n".join(lines), {"page_count": page_count, "source_filename": filename}


def _convert_docx(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    document = DocxDocument(BytesIO(content))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        table_lines: list[str] = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                table_lines.append(" | ".join(values))
        if table_lines:
            lines.append("")
            lines.append(f"## Table {table_index}")
            lines.extend(table_lines)
    return "\n".join(lines), {"source_filename": filename}


def _convert_xlsx(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
    lines: list[str] = []
    sheet_names = list(workbook.sheetnames)
    for sheet_name in sheet_names:
        sheet = workbook[sheet_name]
        rows = list(sheet.iter_rows(values_only=True))
        non_empty_rows = [
            ["" if value is None else str(value).strip() for value in row]
            for row in rows
            if any(value not in (None, "") for value in row)
        ]
        if not non_empty_rows:
            continue
        if lines:
            lines.append("")
        lines.append(f"# Sheet: {sheet_name}")
        header = non_empty_rows[0]
        lines.append(" | ".join(header))
        for row in non_empty_rows[1:]:
            padded = row + [""] * max(0, len(header) - len(row))
            lines.append(" | ".join(padded[: len(header)]))
    return "\n".join(lines), {"sheet_names": sheet_names, "source_filename": filename}


def _convert_csv(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    text = _decode_text_content(content)
    reader = csv.reader(text.splitlines())
    rows = [["" if cell is None else cell.strip() for cell in row] for row in reader]
    rows = [row for row in rows if any(row)]
    lines = [" | ".join(row) for row in rows]
    return "\n".join(lines), {"row_count": len(rows), "source_filename": filename}


def _convert_text(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    return _decode_text_content(content).strip(), {"source_filename": filename}


def _decode_text_content(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError("Unable to decode text document with supported encodings.")


def _convert_with_markitdown(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise DocumentProcessingError(
            "markitdown is required for document conversion. Install with "
            "`pip install markitdown`."
        ) from exc

    converter = MarkItDown()
    stream = BytesIO(content)
    stream.name = Path(filename).name  # type: ignore[attr-defined]
    result = converter.convert_stream(stream, file_extension=Path(filename).suffix)
    text = getattr(result, "text_content", None)
    metadata = getattr(result, "metadata", None)
    return str(text or ""), metadata if isinstance(metadata, dict) else {}


def _document_type(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def _converter_version(package_name: str) -> str:
    try:
        return version(package_name)
    except PackageNotFoundError:
        return "unknown"


_NATIVE_CONVERTERS: dict[str, Any] = {
    "pdf": _convert_pdf,
    "docx": _convert_docx,
    "xlsx": _convert_xlsx,
    "csv": _convert_csv,
    "json": _convert_text,
    "xml": _convert_text,
    "md": _convert_text,
    "markdown": _convert_text,
    "txt": _convert_text,
    "html": _convert_text,
    "htm": _convert_text,
}
