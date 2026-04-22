from __future__ import annotations

import json
import os
import re
from dataclasses import asdict
from functools import lru_cache
from importlib.metadata import PackageNotFoundError, version
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from PIL import Image
from docx import Document as DocxDocument
from openpyxl import load_workbook

from semantic_reasoning_agent.core.config import Settings
from semantic_reasoning_agent.documents.errors import DocumentProcessingError
from semantic_reasoning_agent.documents.models import (
    DocumentIngestionOptions,
    ParsedArtifact,
    ParsedChunk,
    ParsedDocument,
    StructuredExtractionResult,
)


PARSER_NAME = "marker"
_PAGE_MARKER_PATTERN = re.compile(r"^\d+$")
_PAGE_DIVIDER_PATTERN = re.compile(r"^-{8,}$")
_IMAGE_EXTENSIONS = ("png", "jpg", "jpeg", "webp", "gif", "bmp", "tif", "tiff")
_FALLBACK_TYPES = frozenset({"pdf", "docx", "xlsx"})
_WHITESPACE_PATTERN = re.compile(r"\s+")


class MarkerDocumentParser:
    supported_types = (
        "pdf",
        "docx",
        "xlsx",
        "pptx",
        "html",
        "epub",
        *_IMAGE_EXTENSIONS,
    )

    def __init__(self, settings: Settings) -> None:
        self._settings = settings

    def parse(
        self,
        filename: str,
        content: bytes,
        title: str | None = None,
        *,
        options: DocumentIngestionOptions | None = None,
    ) -> ParsedDocument:
        resolved_options = options or DocumentIngestionOptions()
        document_type = _document_type(filename)

        if self._settings.app_env == "test" and document_type in _FALLBACK_TYPES:
            return _parse_with_fallback(document_type, filename, content, title, resolved_options)

        try:
            return self._parse_with_marker(filename, content, title=title, options=resolved_options)
        except ImportError as exc:
            if _should_use_fallback(document_type, resolved_options):
                return _parse_with_fallback(document_type, filename, content, title, resolved_options)
            raise DocumentProcessingError(_marker_dependency_error()) from exc
        except DocumentProcessingError as exc:
            if _marker_runtime_unavailable(exc) and _should_use_fallback(document_type, resolved_options):
                return _parse_with_fallback(document_type, filename, content, title, resolved_options)
            raise
        except Exception as exc:
            raise DocumentProcessingError(str(exc)) from exc

    def extract_structured(
        self,
        filename: str,
        content: bytes,
        *,
        schema_json: dict[str, Any],
        use_llm: bool,
        force_ocr: bool,
        strip_existing_ocr: bool,
        existing_markdown: str | None = None,
    ) -> StructuredExtractionResult:
        if self._settings.app_env == "test":
            markdown = existing_markdown or content.decode("utf-8", errors="ignore")
            return StructuredExtractionResult(
                parser_name=PARSER_NAME,
                parser_version=_marker_version(),
                result={
                    "schema": schema_json,
                    "items": [line.strip() for line in markdown.splitlines() if line.strip()][:5],
                },
                markdown=markdown,
            )

        self._configure_runtime()
        try:
            from marker.config.parser import ConfigParser
            from marker.converters.extraction import ExtractionConverter
        except ImportError as exc:
            raise DocumentProcessingError(_marker_dependency_error()) from exc

        config = self._build_marker_config(
            output_format="json",
            use_llm=use_llm,
            force_ocr=force_ocr,
            strip_existing_ocr=strip_existing_ocr,
            extract_images=False,
            page_schema=schema_json,
            existing_markdown=existing_markdown,
        )
        config_parser = ConfigParser(config)
        converter = ExtractionConverter(
            artifact_dict=_get_marker_models(),
            config=config_parser.generate_config_dict(),
            llm_service=self._resolve_llm_service(config_parser, use_llm),
        )

        with TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / filename
            path.write_bytes(content)
            rendered = converter(str(path))

        payload = _safe_model_dump(rendered)
        return StructuredExtractionResult(
            parser_name=PARSER_NAME,
            parser_version=_marker_version(),
            result=payload,
            markdown=payload.get("original_markdown") or payload.get("markdown"),
        )

    def _parse_with_marker(
        self,
        filename: str,
        content: bytes,
        *,
        title: str | None,
        options: DocumentIngestionOptions,
    ) -> ParsedDocument:
        markdown_result = self._run_marker(filename, content, options=options, output_format="markdown")
        if options.output_format == "chunks":
            # Chunks are derived from markdown; no second Marker pass needed.
            chunk_payload = [asdict(chunk) for chunk in _markdown_to_chunks(markdown_result["text"])]
            requested_result = {
                "text": json.dumps(chunk_payload),
                "extension": "json",
                "images": markdown_result["images"],
                "payload": {"chunks": chunk_payload},
                "metadata": markdown_result["metadata"],
                "page_count": markdown_result["page_count"],
                "warnings": markdown_result["warnings"],
            }
        else:
            requested_result = (
                markdown_result
                if options.output_format == "markdown"
                else self._run_marker(filename, content, options=options, output_format=options.output_format)
            )

        markdown = markdown_result["text"]
        if not markdown.strip():
            raise DocumentProcessingError(f"No extractable text was found in '{filename}'.")

        metadata = markdown_result["metadata"]
        page_count = _page_count_from(metadata) or markdown_result["page_count"]
        quality_flags = _build_quality_flags(metadata, options)
        warnings = tuple(markdown_result["warnings"])

        artifacts = _build_artifacts(
            markdown_result=markdown_result,
            requested_result=requested_result,
            requested_format=options.output_format,
        )

        requested_payload = requested_result["payload"]
        return ParsedDocument(
            document_type=_document_type(filename),
            title=title or Path(filename).stem,
            chunks=tuple(_markdown_to_chunks(markdown)),
            parser_name=PARSER_NAME,
            parser_version=_marker_version(),
            markdown=markdown,
            html=requested_result["text"] if options.output_format == "html" else None,
            json_payload=requested_payload if isinstance(requested_payload, dict) else None,
            metadata=metadata,
            artifacts=tuple(artifacts),
            page_count=page_count,
            quality_flags=quality_flags,
            warnings=warnings,
            ocr_used="ocr_used" in quality_flags,
        )

    def _run_marker(
        self,
        filename: str,
        content: bytes,
        *,
        options: DocumentIngestionOptions,
        output_format: str,
    ) -> dict[str, Any]:
        self._configure_runtime()
        try:
            from marker.config.parser import ConfigParser
            from marker.output import text_from_rendered
        except ImportError as exc:
            raise DocumentProcessingError(_marker_dependency_error()) from exc

        config = self._build_marker_config(
            output_format=output_format,
            use_llm=options.use_llm or (
                options.pdf_mode == "accurate" and self._settings.marker_use_llm_in_accurate
            ),
            force_ocr=options.force_ocr or options.pdf_mode == "accurate",
            strip_existing_ocr=options.strip_existing_ocr,
            extract_images=options.extract_images,
        )
        config_parser = ConfigParser(config)
        converter_cls = config_parser.get_converter_cls()
        converter = converter_cls(
            config=config_parser.generate_config_dict(),
            artifact_dict=_get_marker_models(),
            processor_list=config_parser.get_processors(),
            renderer=config_parser.get_renderer(),
            llm_service=self._resolve_llm_service(config_parser, bool(config.get("use_llm"))),
        )

        with TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / filename
            path.write_bytes(content)
            rendered = converter(str(path))

        text, extension, images = text_from_rendered(rendered)
        payload = _safe_model_dump(rendered)
        metadata = getattr(rendered, "metadata", None) or payload.get("metadata") or {}
        return {
            "text": text if isinstance(text, str) else "",
            "extension": extension,
            "images": images or {},
            "payload": payload,
            "metadata": metadata if isinstance(metadata, dict) else {},
            "page_count": getattr(converter, "page_count", None),
            "warnings": _marker_warnings(output_format, options.output_format),
        }

    def _build_marker_config(
        self,
        *,
        output_format: str,
        use_llm: bool,
        force_ocr: bool,
        strip_existing_ocr: bool,
        extract_images: bool,
        page_schema: dict[str, Any] | None = None,
        existing_markdown: str | None = None,
    ) -> dict[str, Any]:
        config: dict[str, Any] = {
            "output_format": output_format,
            "paginate_output": output_format == "markdown",
            "disable_image_extraction": not extract_images,
            "force_ocr": force_ocr,
            "strip_existing_ocr": strip_existing_ocr,
            "use_llm": use_llm,
        }
        if page_schema is not None:
            config["page_schema"] = page_schema
        if existing_markdown:
            config["existing_markdown"] = existing_markdown
        if self._settings.marker_max_pages_per_doc:
            config["page_range"] = list(range(1, self._settings.marker_max_pages_per_doc + 1))
        if self._settings.marker_torch_device:
            config["TORCH_DEVICE"] = self._settings.marker_torch_device

        llm_config = self._build_llm_config()
        config.update(llm_config)
        return config

    def _build_llm_config(self) -> dict[str, Any]:
        config: dict[str, Any] = {}
        if self._settings.google_api_key:
            config["gemini_api_key"] = self._settings.google_api_key
        if self._settings.openai_api_key:
            config["openai_api_key"] = self._settings.openai_api_key
        if self._settings.openai_base_url:
            config["openai_base_url"] = self._settings.openai_base_url
        if self._settings.anthropic_api_key:
            config["claude_api_key"] = self._settings.anthropic_api_key
        return config

    def _resolve_llm_service(self, config_parser: Any, use_llm: bool) -> str | None:
        if not use_llm:
            return None
        if self._settings.google_api_key:
            return "marker.services.gemini.GoogleGeminiService"
        if self._settings.openai_api_key:
            return "marker.services.openai.OpenAIService"
        if self._settings.anthropic_api_key:
            return "marker.services.claude.ClaudeService"
        return config_parser.get_llm_service()

    def _configure_runtime(self) -> None:
        cache_dir = Path(self._settings.marker_model_cache_dir)
        cache_dir.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault("MODEL_CACHE_DIR", str(cache_dir))
        if self._settings.marker_torch_device:
            os.environ.setdefault("TORCH_DEVICE", self._settings.marker_torch_device)
        try:
            from surya.settings import settings as surya_settings

            surya_settings.MODEL_CACHE_DIR = str(cache_dir)
            if self._settings.marker_torch_device:
                surya_settings.TORCH_DEVICE = self._settings.marker_torch_device
        except Exception:
            pass


@lru_cache(maxsize=1)
def _get_marker_models():
    from marker.models import create_model_dict

    return create_model_dict()


@lru_cache(maxsize=1)
def _marker_version() -> str:
    try:
        return f"marker-v{version('marker-pdf')}"
    except PackageNotFoundError:
        return "marker-vunknown"


def _marker_dependency_error() -> str:
    return "marker-pdf[full] is required for Marker-based document ingestion."


def _marker_runtime_unavailable(exc: DocumentProcessingError) -> bool:
    return _marker_dependency_error() in str(exc)


def _should_use_fallback(
    document_type: str,
    options: DocumentIngestionOptions,
) -> bool:
    if document_type not in _FALLBACK_TYPES:
        return False
    if document_type == "pdf" and options.pdf_mode == "accurate":
        return False
    return True


def _marker_warnings(rendered_format: str, requested_format: str) -> tuple[str, ...]:
    if rendered_format != requested_format:
        return (f"marker_rendered_as_{rendered_format}_for_chunking",)
    return ()


def _safe_model_dump(rendered: Any) -> dict[str, Any]:
    try:
        return rendered.model_dump(mode="json")
    except Exception:
        return json.loads(rendered.model_dump_json())


def _build_artifacts(
    *,
    markdown_result: dict[str, Any],
    requested_result: dict[str, Any],
    requested_format: str,
) -> list[ParsedArtifact]:
    artifacts = list(_build_image_artifacts(markdown_result["images"]))
    artifacts.append(
        ParsedArtifact(
            artifact_type="markdown",
            name="document.md",
            content=markdown_result["text"].encode("utf-8"),
            content_type="text/markdown",
        )
    )
    artifacts.append(
        ParsedArtifact(
            artifact_type="marker_json",
            name="document.json",
            content=json.dumps(markdown_result["payload"], ensure_ascii=False).encode("utf-8"),
            content_type="application/json",
        )
    )

    if requested_format == "html":
        artifacts.append(
            ParsedArtifact(
                artifact_type="html",
                name="document.html",
                content=requested_result["text"].encode("utf-8"),
                content_type="text/html",
            )
        )
    elif requested_format in {"json", "chunks"}:
        artifacts.append(
            ParsedArtifact(
                artifact_type=requested_format,
                name=f"document.{requested_result['extension']}",
                content=requested_result["text"].encode("utf-8"),
                content_type="application/json",
            )
        )

    return artifacts


def _page_count_from(metadata: dict[str, Any]) -> int | None:
    page_stats = metadata.get("page_stats")
    if isinstance(page_stats, list) and page_stats:
        return len(page_stats)
    return None


def _build_quality_flags(
    metadata: dict[str, Any],
    options: DocumentIngestionOptions,
) -> tuple[str, ...]:
    quality_flags: list[str] = []
    if options.force_ocr or options.pdf_mode == "accurate":
        quality_flags.append("ocr_forced")

    page_stats = metadata.get("page_stats")
    if isinstance(page_stats, list):
        for page in page_stats:
            if not isinstance(page, dict):
                continue
            if page.get("text_extraction_method") == "surya":
                quality_flags.append("ocr_used")
                break

    return tuple(dict.fromkeys(quality_flags))


def _build_image_artifacts(images: dict[str, Image.Image]) -> list[ParsedArtifact]:
    artifacts: list[ParsedArtifact] = []
    for index, (image_id, image) in enumerate((images or {}).items(), start=1):
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        artifacts.append(
            ParsedArtifact(
                artifact_type="image",
                name=f"image-{index}.png",
                content=buffer.getvalue(),
                content_type="image/png",
                metadata={"marker_image_id": image_id},
            )
        )
    return artifacts


def _markdown_to_chunks(markdown: str) -> list[ParsedChunk]:
    heading_stack: list[str] = []
    blocks: list[ParsedChunk] = []
    buffer: list[str] = []
    current_page: int | None = None
    lines = markdown.splitlines()
    index = 0

    def flush() -> None:
        nonlocal buffer
        text = "\n".join(line for line in buffer if line.strip()).strip()
        buffer = []
        if not text:
            return
        section_title = heading_stack[-1] if heading_stack else None
        blocks.append(
            ParsedChunk(
                text=text,
                chunk_index=len(blocks),
                page_number=current_page,
                section_title=section_title,
                heading_path=" > ".join(heading_stack) if heading_stack else None,
            )
        )

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""

        if _PAGE_MARKER_PATTERN.match(stripped) and _PAGE_DIVIDER_PATTERN.match(next_line):
            flush()
            current_page = int(stripped) + 1
            index += 2
            continue
        if stripped.startswith("#"):
            flush()
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if title:
                heading_stack = heading_stack[: level - 1]
                heading_stack.append(title)
            index += 1
            continue
        if stripped.startswith("![](") or stripped.startswith("<img"):
            index += 1
            continue
        if not stripped and buffer:
            buffer.append("")
            index += 1
            continue
        if stripped:
            buffer.append(stripped)
        index += 1

    flush()
    return blocks or [ParsedChunk(text=markdown.strip(), chunk_index=0, page_number=current_page)]


def _document_type(filename: str) -> str:
    return Path(filename).suffix.lower().lstrip(".")


def _parse_with_fallback(
    document_type: str,
    filename: str,
    content: bytes,
    title: str | None,
    options: DocumentIngestionOptions,
) -> ParsedDocument:
    if document_type == "pdf":
        return _parse_pdf_fallback(filename, content, title, options)
    if document_type == "docx":
        return _parse_docx_fallback(filename, content, title)
    if document_type == "xlsx":
        return _parse_xlsx_fallback(filename, content, title)
    raise DocumentProcessingError(f"Unsupported fallback parser for '{document_type}'.")


def _parse_pdf_fallback(
    filename: str,
    content: bytes,
    title: str | None,
    options: DocumentIngestionOptions,
) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # noqa: BLE001
        raise DocumentProcessingError(_marker_dependency_error()) from exc

    reader = PdfReader(BytesIO(content))
    chunks: list[ParsedChunk] = []
    markdown_lines: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        markdown_lines.extend([str(page_index - 1), "--------", text, ""])
        chunks.append(
            ParsedChunk(
                text=text,
                chunk_index=len(chunks),
                page_number=page_index,
                section_title=f"Page {page_index}",
                heading_path=f"Page {page_index}",
            )
        )
    if not chunks:
        raise DocumentProcessingError(f"No extractable text was found in '{filename}'.")
    markdown = "\n".join(markdown_lines).strip()
    return ParsedDocument(
        document_type="pdf",
        title=title or Path(filename).stem,
        chunks=tuple(chunks),
        parser_name="pypdf",
        parser_version="pypdf-fallback-v1",
        markdown=markdown,
        page_count=len(chunks),
        warnings=("marker_runtime_unavailable",),
    )


def _parse_docx_fallback(
    filename: str,
    content: bytes,
    title: str | None,
) -> ParsedDocument:
    document = DocxDocument(BytesIO(content))
    chunks: list[ParsedChunk] = []
    heading_stack: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraphs() -> None:
        if not paragraph_buffer:
            return
        heading_path = " > ".join(heading_stack) if heading_stack else "Document"
        chunks.append(
            ParsedChunk(
                text="\n".join(paragraph_buffer),
                chunk_index=len(chunks),
                section_title=heading_stack[-1] if heading_stack else "Document",
                heading_path=heading_path,
            )
        )
        paragraph_buffer.clear()

    for paragraph in document.paragraphs:
        text = _normalize_inline_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("Heading"):
            flush_paragraphs()
            level = _extract_heading_level(style_name)
            heading_stack[:] = heading_stack[: level - 1]
            heading_stack.append(text)
            continue
        paragraph_buffer.append(text)

    flush_paragraphs()

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [_normalize_inline_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if not rows:
            continue
        heading_path = " > ".join(heading_stack) if heading_stack else "Document"
        chunks.append(
            ParsedChunk(
                text="\n".join(rows),
                chunk_index=len(chunks),
                section_title=heading_stack[-1] if heading_stack else "Document",
                heading_path=heading_path,
                table_index=table_index,
            )
        )

    if not chunks:
        raise DocumentProcessingError(f"No extractable content was found in '{filename}'.")

    return ParsedDocument(
        document_type="docx",
        title=title or Path(filename).stem,
        chunks=tuple(chunks),
        parser_name="python-docx",
        parser_version="docx-structured-v2",
        warnings=("marker_runtime_unavailable",),
    )


def _parse_xlsx_fallback(
    filename: str,
    content: bytes,
    title: str | None,
) -> ParsedDocument:
    workbook = load_workbook(BytesIO(content), data_only=True)
    chunks: list[ParsedChunk] = []
    sheet_names: list[str] = []

    for sheet in workbook.worksheets:
        sheet_names.append(sheet.title)
        non_empty_rows: list[tuple[int, list[str]]] = []
        for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [_stringify_cell(cell) for cell in row]
            if any(values):
                non_empty_rows.append((row_number, values))

        if not non_empty_rows:
            continue

        header_row_number, header_values = non_empty_rows[0]
        headers = [
            header if header else f"column_{column_index}"
            for column_index, header in enumerate(header_values, start=1)
        ]
        chunks.append(
            ParsedChunk(
                text=f"Sheet: {sheet.title}\nColumns: {', '.join(headers)}",
                chunk_index=len(chunks),
                section_title=sheet.title,
                sheet_name=sheet.title,
                row_start=header_row_number,
                row_end=header_row_number,
                column_headers=tuple(headers),
                detected_table_id=f"{sheet.title}:schema",
            )
        )

        data_rows = non_empty_rows[1:]
        for offset in range(0, len(data_rows), 5):
            group = data_rows[offset : offset + 5]
            lines: list[str] = []
            for row_number, values in group:
                pairs = [f"{header}: {value}" for header, value in zip(headers, values) if value]
                if pairs:
                    lines.append(f"Row {row_number}: " + "; ".join(pairs))
            if not lines:
                continue
            row_start = group[0][0]
            row_end = group[-1][0]
            chunks.append(
                ParsedChunk(
                    text=f"Sheet: {sheet.title}\nColumns: {', '.join(headers)}\n" + "\n".join(lines),
                    chunk_index=len(chunks),
                    section_title=sheet.title,
                    sheet_name=sheet.title,
                    row_start=row_start,
                    row_end=row_end,
                    column_headers=tuple(headers),
                    detected_table_id=f"{sheet.title}:{row_start}-{row_end}",
                )
            )

    if not chunks:
        raise DocumentProcessingError(f"No extractable worksheet content was found in '{filename}'.")

    return ParsedDocument(
        document_type="xlsx",
        title=title or Path(filename).stem,
        chunks=tuple(chunks),
        parser_name="openpyxl",
        parser_version="xlsx-structured-v2",
        sheet_names=tuple(sheet_names),
        warnings=("marker_runtime_unavailable",),
    )
def _normalize_inline_text(value: str) -> str:
    return _WHITESPACE_PATTERN.sub(" ", value or "").strip()


def _extract_heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)$", style_name)
    if match:
        return max(1, int(match.group(1)))
    return 1


def _stringify_cell(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()
