from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from semantic_reasoning_agent.documents.errors import DocumentProcessingError
from semantic_reasoning_agent.documents.models import DocumentIngestionOptions, ParsedChunk, ParsedDocument


PARSER_NAME = "python-docx"
PARSER_VERSION = "docx-structured-v2"


class DocxParser:
    supported_types = ("docx",)

    def parse(
        self,
        filename: str,
        content: bytes,
        title: str | None = None,
        *,
        options: DocumentIngestionOptions | None = None,
    ) -> ParsedDocument:
        del options
        document = DocxDocument(BytesIO(content))
        chunks: list[ParsedChunk] = []
        heading_stack: list[str] = []
        paragraph_buffer: list[str] = []

        def flush_paragraphs() -> None:
            if not paragraph_buffer:
                return
            heading_path = " > ".join(heading_stack) if heading_stack else "Document"
            chunks.append(
                ParsedChunk(
                    text="\n".join(paragraph_buffer),
                    chunk_index=len(chunks),
                    section_title=heading_stack[-1] if heading_stack else "Document",
                    heading_path=heading_path,
                )
            )
            paragraph_buffer.clear()

        for paragraph in document.paragraphs:
            text = _normalize_inline_text(paragraph.text)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
            if style_name.startswith("Heading"):
                flush_paragraphs()
                level = _extract_heading_level(style_name)
                heading_stack[:] = heading_stack[: level - 1]
                heading_stack.append(text)
                continue
            paragraph_buffer.append(text)

        flush_paragraphs()

        for table_index, table in enumerate(document.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [_normalize_inline_text(cell.text) for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            heading_path = " > ".join(heading_stack) if heading_stack else "Document"
            chunks.append(
                ParsedChunk(
                    text="\n".join(rows),
                    chunk_index=len(chunks),
                    section_title=heading_stack[-1] if heading_stack else "Document",
                    heading_path=heading_path,
                    table_index=table_index,
                )
            )

        if not chunks:
            raise DocumentProcessingError(f"No extractable content was found in '{filename}'.")

        return ParsedDocument(
            document_type="docx",
            title=title or Path(filename).stem,
            chunks=tuple(chunks),
            parser_name=PARSER_NAME,
            parser_version=PARSER_VERSION,
        )


def _normalize_inline_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _extract_heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)$", style_name)
    if match:
        return max(1, int(match.group(1)))
    return 1
