import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from semantic_reasoning_agent.infrastructure.parsers.models import ParsedChunk, ParsedDocument


PARSER_VERSION = "local-structured-v1"


class UnsupportedDocumentTypeError(ValueError):
    """Raised when the upload format is not supported."""


def parse_document(filename: str, content: bytes, title: str | None = None) -> ParsedDocument:
    document_type = Path(filename).suffix.lower().lstrip(".")
    if document_type == "pdf":
        return _parse_pdf(filename, content, title)
    if document_type == "docx":
        return _parse_docx(filename, content, title)
    if document_type == "xlsx":
        return _parse_xlsx(filename, content, title)
    raise UnsupportedDocumentTypeError(
        f"Unsupported document type '{document_type}'. Phase 2 currently supports pdf, docx, and xlsx."
    )


def _parse_pdf(filename: str, content: bytes, title: str | None) -> ParsedDocument:
    reader = PdfReader(BytesIO(content))
    chunks: list[ParsedChunk] = []
    for page_index, page in enumerate(reader.pages, start=1):
        extracted = _normalize_block_text(page.extract_text() or "")
        if not extracted:
            continue
        lines = [line.strip() for line in extracted.splitlines() if line.strip()]
        section_title = lines[0] if lines else None
        chunks.append(
            ParsedChunk(
                text=extracted,
                chunk_index=len(chunks),
                page_number=page_index,
                section_title=section_title,
            )
        )

    if not chunks:
        raise ValueError(f"No extractable text was found in '{filename}'.")

    return ParsedDocument(
        document_type="pdf",
        title=title or Path(filename).stem,
        chunks=chunks,
        parser_version=PARSER_VERSION,
    )


def _parse_docx(filename: str, content: bytes, title: str | None) -> ParsedDocument:
    document = DocxDocument(BytesIO(content))
    chunks: list[ParsedChunk] = []
    heading_stack: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraphs() -> None:
        if not paragraph_buffer:
            return
        heading_path = " > ".join(heading_stack) if heading_stack else "Document"
        chunks.append(
            ParsedChunk(
                text="\n".join(paragraph_buffer),
                chunk_index=len(chunks),
                section_title=heading_stack[-1] if heading_stack else "Document",
                heading_path=heading_path,
            )
        )
        paragraph_buffer.clear()

    for paragraph in document.paragraphs:
        text = _normalize_inline_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
        if style_name.startswith("Heading"):
            flush_paragraphs()
            level = _extract_heading_level(style_name)
            heading_stack[:] = heading_stack[: level - 1]
            heading_stack.append(text)
            continue
        paragraph_buffer.append(text)

    flush_paragraphs()

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [_normalize_inline_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        if not rows:
            continue
        heading_path = " > ".join(heading_stack) if heading_stack else "Document"
        chunks.append(
            ParsedChunk(
                text="\n".join(rows),
                chunk_index=len(chunks),
                section_title=heading_stack[-1] if heading_stack else "Document",
                heading_path=heading_path,
                table_index=table_index,
            )
        )

    if not chunks:
        raise ValueError(f"No extractable content was found in '{filename}'.")

    return ParsedDocument(
        document_type="docx",
        title=title or Path(filename).stem,
        chunks=chunks,
        parser_version=PARSER_VERSION,
    )


def _parse_xlsx(filename: str, content: bytes, title: str | None) -> ParsedDocument:
    workbook = load_workbook(BytesIO(content), data_only=True)
    chunks: list[ParsedChunk] = []
    for sheet in workbook.worksheets:
        non_empty_rows: list[tuple[int, list[str]]] = []
        for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [_stringify_cell(cell) for cell in row]
            if any(values):
                non_empty_rows.append((row_number, values))

        if not non_empty_rows:
            continue

        header_row_number, header_values = non_empty_rows[0]
        headers = [
            header if header else f"column_{column_index}"
            for column_index, header in enumerate(header_values, start=1)
        ]
        chunks.append(
            ParsedChunk(
                text=f"Sheet: {sheet.title}\nColumns: {', '.join(headers)}",
                chunk_index=len(chunks),
                section_title=sheet.title,
                sheet_name=sheet.title,
                row_start=header_row_number,
                row_end=header_row_number,
                column_headers=headers,
                detected_table_id=f"{sheet.title}:schema",
            )
        )

        data_rows = non_empty_rows[1:]
        for offset in range(0, len(data_rows), 5):
            group = data_rows[offset : offset + 5]
            lines: list[str] = []
            for row_number, values in group:
                pairs = [
                    f"{header}: {value}"
                    for header, value in zip(headers, values)
                    if value
                ]
                if pairs:
                    lines.append(f"Row {row_number}: " + "; ".join(pairs))
            if not lines:
                continue
            row_start = group[0][0]
            row_end = group[-1][0]
            chunks.append(
                ParsedChunk(
                    text=(
                        f"Sheet: {sheet.title}\n"
                        f"Columns: {', '.join(headers)}\n"
                        + "\n".join(lines)
                    ),
                    chunk_index=len(chunks),
                    section_title=sheet.title,
                    sheet_name=sheet.title,
                    row_start=row_start,
                    row_end=row_end,
                    column_headers=headers,
                    detected_table_id=f"{sheet.title}:{row_start}-{row_end}",
                )
            )

    if not chunks:
        raise ValueError(f"No extractable worksheet content was found in '{filename}'.")

    return ParsedDocument(
        document_type="xlsx",
        title=title or Path(filename).stem,
        chunks=chunks,
        parser_version=PARSER_VERSION,
    )


def _normalize_block_text(value: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line)


def _normalize_inline_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _extract_heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)$", style_name)
    if match:
        return max(1, int(match.group(1)))
    return 1


def _stringify_cell(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()
